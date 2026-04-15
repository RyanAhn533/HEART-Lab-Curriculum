"""
HEART Lab AI Curriculum — DOCX 생성 스크립트
CURRICULUM_MASTER.md 내용을 정식 문서 형태로 변환
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '맑은 고딕'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ── 헬퍼 함수 ──
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9.5)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ══════════════════════════════════════════
# 표지
# ══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HEART Lab\nAI Education Curriculum')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.bold = True

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('세종대학교 HEART Lab\n신입 연구원 AI 교육 과정 설계 문서')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('작성: 안준영  |  HEART Lab  |  2026.04')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ══════════════════════════════════════════
# 목차
# ══════════════════════════════════════════
doc.add_heading('목차', level=1)
toc_items = [
    '1. 교육 이념',
    '2. 교육 원칙',
    '3. 이론 학습 원칙',
    '4. 운영 방식',
    '5. 전체 로드맵 (Phase 0~10)',
    '6. 변화 포인트 추적',
    '7. 이론 매핑 전체표',
    '8. 세계 트렌드와 HEART Lab 연결',
    '9. 폴더 구조',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════
# 1. 교육 이념
# ══════════════════════════════════════════
doc.add_heading('1. 교육 이념', level=1)

doc.add_heading('한 줄 정의', level=2)
p = doc.add_paragraph()
run = p.add_run('"AI로 문제를 풀 수 있는 사람을 만든다"')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_heading('왜 이렇게 가르치는가', level=2)

doc.add_heading('시장 현실', level=3)
bullets = [
    'AI 경력직 요구 비율: 54% → 80.6%',
    '2026년 신입 채용 비율: 12.4%',
    'SW 신입 비중 2년 변화: 53.5% → 37.4%',
    '기업은 바로 투입 가능한 사람만 원한다',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('바로 투입이란', level=3)
bullets = [
    '"YOLO 써봤습니다" X',
    '"이 문제에 왜 YOLO인지 설명하고, 설계하고, 구현할 수 있습니다" O',
    '코딩은 기본, 뭘 만들지 아는 것까지 되어야 함',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Top-Down Thinking', level=3)
bullets = [
    'First Principles — 근본으로 돌아가서 문제를 재정의',
    '교육도 Top-Down: 일단 돌려보고 → 왜 그런지 파고든다',
    '수학부터 X, 결과부터 보고 필요한 이론을 역추적',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('3단계 성장 모델', level=2)
add_table(doc,
    ['단계', '목표', '기준'],
    [
        ['1단계: 문제 해결 사고', '"이 문제 어떻게 풀래?" 바로 떠오를 때까지 반복', 'Phase 2~3'],
        ['2단계: 모델을 도구로', '"이 문제에 왜 이 모델인지" 설명할 수 있는 것', 'Phase 4~8'],
        ['3단계: 서비스 구현', '챗봇, 실시간 탐지 — 혼자 만들 수 있는 것', 'Phase 9~10'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 2. 교육 원칙
# ══════════════════════════════════════════
doc.add_heading('2. 교육 원칙', level=1)

doc.add_heading('코드 5대 원칙', level=2)
add_table(doc,
    ['원칙', '설명'],
    [
        ['구조 통일', '모든 코드가 #0→#1→#2→#3→#4 동일 뼈대'],
        ['하나만 바뀐다', '이전 코드에서 변경된 건 항상 1개 섹션'],
        ['같은 데이터, 다른 기법', '5개 기본 데이터셋을 끝까지 반복 적용'],
        ['간결하게', 'TF2/Keras의 간결한 API만, 복잡한 건 안 씀'],
        ['60~120줄', '파일당 최대 120줄, 그 이상이면 분리'],
    ]
)

doc.add_heading('코드 뼈대 (모든 파일 공통)', level=2)
code = """# ============================================
# [번호] [주제] — [데이터셋]
# 이전과 차이: [한 줄로 뭐가 바뀌었는지]
#
# 왜 배우는가:
#   [이 기법이 왜 필요한지]
#
# ▶ 보고 오기: [필수 영상/블로그]
# ▶ 나중에 읽기: [논문]
# ============================================

#0. 라이브러리
#1. 데이터
#2. 데이터 전처리 및 분할
#3. 모델
#4. 평가"""
add_code_block(doc, code)

doc.add_heading('기술 스택 규칙', level=2)
add_table(doc,
    ['구간', '프레임워크', '이유'],
    [
        ['Phase 0~8', 'TensorFlow2/Keras만', '간결, 선언적, 결과 빨리 봄'],
        ['Phase 9', 'TF2 + sklearn', 'ML 기법은 sklearn이 표준'],
        ['Phase 10', 'PyTorch', '내부 동작 이해, 논문 재현'],
    ]
)

doc.add_heading('5개 기본 데이터셋', level=2)
add_table(doc,
    ['데이터셋', '유형', '특성 수', '용도'],
    [
        ['Boston Housing', '회귀', '13', '기본 회귀'],
        ['Diabetes', '회귀', '10', '의료 데이터'],
        ['Iris', '다중분류', '4', '분류 입문'],
        ['Breast Cancer', '이진분류', '30', '이진분류'],
        ['MNIST', '이미지분류', '28x28', 'CNN 이후'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 3. 이론 학습 원칙
# ══════════════════════════════════════════
doc.add_heading('3. 이론 학습 원칙', level=1)

doc.add_heading('학습 순서', level=2)
p = doc.add_paragraph()
run = p.add_run('영상 (직관)  →  블로그 리뷰 (정리)  →  논문 (원본)')
run.bold = True
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('절대 논문부터 읽지 마라.')
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True

doc.add_heading('리소스 우선순위', level=2)

doc.add_heading('1순위: 3Blue1Brown (영상)', level=3)
bullets = [
    '시각적 직관 최강. 한글자막 있음.',
    '이해 안 되면 3Blue1Brown부터 찾아라.',
    '핵심 시리즈: Essence of Linear Algebra, Neural Networks, Calculus, Transformers',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('2순위: 구글 검색 "[주제] 논문 리뷰"', level=3)
bullets = [
    '"LSTM 논문 리뷰", "ResNet 논문 리뷰" 이렇게 검색',
    '한국어 블로그 2~3개 읽으면 핵심 파악 가능',
    '논문 원문 읽기 전에 리뷰 먼저 보면 맥락이 잡힘',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('3순위: 특정 추천 블로그/채널', level=3)
bullets = [
    'Transformer 계열: codingopera.tistory.com/43',
    'StatQuest — 통계/ML 기초 (영어)',
    '혁펜하임 — 한국어 DL 전반',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('필수 논문 (읽는 시점)', level=2)
add_table(doc,
    ['시점', '논문', '왜 읽는가'],
    [
        ['Phase 5 후', 'LeNet-5 (LeCun, 1998)', 'CNN의 시작'],
        ['Phase 5 후', 'AlexNet (Krizhevsky, 2012)', '딥러닝 부활'],
        ['Phase 6 후', 'LSTM (Hochreiter, 1997)', '시계열의 핵심'],
        ['Phase 7 후', 'Word2Vec (Mikolov, 2013)', '임베딩의 시작'],
        ['Phase 8', 'VGGNet (Simonyan, 2014)', '전이학습 기본'],
        ['Phase 8', 'ResNet (He, 2015)', 'Skip Connection'],
        ['Phase 10', 'Attention Is All You Need (Vaswani, 2017)', '현대 AI의 기반'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 4. 운영 방식
# ══════════════════════════════════════════
doc.add_heading('4. 운영 방식', level=1)

doc.add_heading('단일 트랙', level=2)
p = doc.add_paragraph()
run = p.add_run('학부/석사 구분 없음. 전부 같은 트랙. 아는 사람은 스킵. 모르면 처음부터.')
run.bold = True

doc.add_heading('진입 방식', level=2)
add_table(doc,
    ['진단 결과', '시작 지점'],
    [
        ['전부 모름', 'Phase 0부터'],
        ['Phase 2까지 알면', 'Phase 3부터'],
        ['Phase 6까지 알면', 'Phase 7부터'],
        ['전부 알면', '논문 재현 프로젝트만'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('스킵 기준: 해당 Phase 코드를 보고 30분 안에 혼자 짤 수 있으면 넘어감.')
run.italic = True

doc.add_heading('평가', level=2)
add_table(doc,
    ['시점', '평가', '범위'],
    [
        ['Phase 3 끝', '체크포인트 시험 1', '기초 회귀/분류 체화'],
        ['Phase 6 끝', '체크포인트 시험 2', 'CNN/RNN 체화'],
        ['Phase 10 끝', '최종 발표', '논문 재현 프로젝트'],
        ['매월', '월간 발표', '학습 현황 + 결과물'],
    ]
)

doc.add_heading('속도', level=2)
add_table(doc,
    ['유형', '예상 기간'],
    [
        ['빠른 사람', '8주'],
        ['보통', '12~16주'],
        ['느린 사람', '20주'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 5. 전체 로드맵
# ══════════════════════════════════════════
doc.add_heading('5. 전체 로드맵', level=1)

phases = [
    {
        'title': 'Phase 0: 왜 하는가',
        'desc': '동기부여 + 큰 그림',
        'items': [
            ('0-1', 'AI 취업시장 현실', '경력직 80%, 신입 12.4%'),
            ('0-2', '바로 투입 = 문제 해결 능력', '코딩 + 뭘 만들지 아는 것'),
            ('0-3', '세계 트렌드', 'NVIDIA, Affectiva, LLM, 멀티모달'),
            ('0-4', 'HEART Lab 과제', '현대차, 두산, 토론토대'),
            ('0-5', 'Top-Down Thinking', 'First Principles'),
            ('0-6', '커리큘럼 소개', '#0→#1→#2→#3→#4 구조'),
        ],
        'ref': '3Blue1Brown "But what is a neural network?"',
    },
    {
        'title': 'Phase 1: 통계/수학 기초',
        'desc': '코드에서 만날 개념의 뿌리 (수학 수업 아님)',
        'items': [
            ('1-1', '데이터를 보는 눈', '평균, 분산, 표준편차 → StandardScaler'),
            ('1-2', '상관관계', '상관계수, scatter → Feature Importance'),
            ('1-3', '확률', '조건부확률, Bayes → Softmax 출력'),
            ('1-4', '손실함수 직관', 'MSE, CE → model.compile(loss=...)'),
            ('1-5', '경사하강법 직관', '산에서 내려가기 → optimizer'),
        ],
        'ref': '3Blue1Brown "Gradient descent"',
    },
    {
        'title': 'Phase 2: 돌려봐 (TF2/Keras)',
        'desc': '#0→#1→#2→#3→#4 뼈대 익히기',
        'items': [
            ('02', '기본 회귀 (Sequential + Dense)', '#0~#4 뼈대 자체'),
            ('03', '다중입력 + 깊은 네트워크', '#3에 레이어 추가'),
            ('04', 'train/test 분할 + R2', '#2에 split 추가'),
            ('05', '실제 데이터셋 5종 적용', '#1에 실제 데이터'),
        ],
        'ref': '3Blue1Brown "Neural Networks" Ch.1',
    },
    {
        'title': 'Phase 3: 왜 안 되지? (TF2/Keras)',
        'desc': '각 섹션에 1개씩 추가',
        'items': [
            ('06', 'validation + EarlyStopping', '#3에 콜백 추가'),
            ('07', '이진분류 (Sigmoid + BCE)', '#3에 sigmoid, loss 변경'),
            ('08', '다중분류 (Softmax + OneHot)', '#3에 softmax, #2에 인코딩'),
            ('09', 'Scaler (MinMax, Standard)', '#2에 Scaler 추가'),
        ],
        'ref': '3Blue1Brown "Backpropagation" Ch.3~4',
        'checkpoint': '체크포인트 시험 1',
    },
    {
        'title': 'Phase 4: 저장하고 구조 바꾸기 (TF2/Keras)',
        'desc': '#3 구조 변경',
        'items': [
            ('10', '모델 저장/로드 + Checkpoint', '#3 뒤에 save 추가'),
            ('11', 'Dropout', '#3에 Dropout 레이어'),
            ('12', '함수형 API (간단히)', '#3 구조 변경'),
        ],
        'ref': 'StatQuest "Dropout"',
    },
    {
        'title': 'Phase 5: 이미지 (TF2/Keras)',
        'desc': '#1, #2 변경 (이미지 데이터)',
        'items': [
            ('13', 'CNN 기초 (Conv2D + MaxPool)', '#1 이미지, #2 reshape, #3 Conv2D'),
            ('14', 'DNN vs CNN 비교', '#3만 교체해서 비교'),
            ('15', '데이터 증강', '#2에 ImageDataGenerator'),
        ],
        'ref': '3Blue1Brown "But what is a convolution?"',
    },
    {
        'title': 'Phase 6: 시계열 (TF2/Keras)',
        'desc': '#2 변경 (3D reshape)',
        'items': [
            ('16', 'SimpleRNN', '#2에 3D reshape, #3에 RNN'),
            ('17', 'LSTM', '#3에 RNN → LSTM 교체'),
            ('18', 'Bidirectional + Conv1D', '#3에 Bidirectional 래핑'),
        ],
        'ref': '3Blue1Brown "Recurrent Neural Networks"',
        'checkpoint': '체크포인트 시험 2',
    },
    {
        'title': 'Phase 7: NLP (TF2/Keras)',
        'desc': '#1, #2 변경 (텍스트 데이터)',
        'items': [
            ('19', 'Tokenizer + Embedding', '#1 텍스트, #2 토큰화, #3 Embedding'),
            ('20', '감정분류 프로젝트', '전체 종합 적용'),
        ],
        'ref': '3Blue1Brown "Word2Vec"',
    },
    {
        'title': 'Phase 8: 성능 올리기 (TF2 + sklearn)',
        'desc': '#3 최적화',
        'items': [
            ('21', '옵티마이저 + ReduceLROnPlateau', '#3에 콜백 추가'),
            ('22', '하이퍼파라미터 튜닝', '#3을 SearchCV로 래핑'),
            ('23', '전이학습 (VGG16, ResNet, GAP)', '#3에 사전학습 모델'),
        ],
        'ref': '구글 "VGGNet 논문 리뷰", "ResNet 논문 리뷰"',
    },
    {
        'title': 'Phase 9: ML 기본기 (sklearn)',
        'desc': 'sklearn 모델',
        'items': [
            ('24', 'PCA + KFold', '#2에 PCA, #4에 cross_val'),
            ('25', 'GridSearch + Feature Importance', '#3을 GridSearchCV로 래핑'),
            ('26', '앙상블 (Bagging, Voting, Stacking)', '#3에 앙상블 모델'),
            ('27', 'SMOTE, PolynomialFeatures', '#2에 SMOTE/Poly 추가'),
        ],
        'ref': '3Blue1Brown "Essence of Linear Algebra" 전편',
    },
    {
        'title': 'Phase 10: PyTorch',
        'desc': '같은 것을 PyTorch로 재구현',
        'items': [
            ('28', 'Keras → Torch 재구현', '같은 #1~#4, 다른 프레임워크'),
            ('29', 'DataLoader, nn.Module', '#1에 DataLoader'),
            ('30', 'CNN + RNN in Torch', '#3에 Conv2d, nn.LSTM'),
            ('31', '논문 재현 프로젝트', 'Transformer, Attention'),
        ],
        'ref': '3Blue1Brown "Transformers", codingopera.tistory.com/43',
        'checkpoint': '최종 발표 → 과제 투입',
    },
]

for phase in phases:
    doc.add_heading(phase['title'], level=2)
    p = doc.add_paragraph()
    run = p.add_run(phase['desc'])
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_table(doc,
        ['번호', '주제', '변화 포인트'],
        [(i[0], i[1], i[2]) for i in phase['items']]
    )

    p = doc.add_paragraph()
    run = p.add_run('▶ 보고 오기: ')
    run.bold = True
    run.font.size = Pt(9.5)
    run = p.add_run(phase['ref'])
    run.font.size = Pt(9.5)

    if 'checkpoint' in phase:
        p = doc.add_paragraph()
        run = p.add_run(f'★ {phase["checkpoint"]}')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_page_break()

# ══════════════════════════════════════════
# 7. 이론 매핑 전체표
# ══════════════════════════════════════════
doc.add_heading('6. 이론 매핑 전체표', level=1)
add_table(doc,
    ['Phase', '코드에서 만나는 것', '보고 오기 (필수)', '나중에 읽기'],
    [
        ['0', '—', '3B1B "Neural Network"', '—'],
        ['1', 'Scaler, Loss, Optimizer', '3B1B "Gradient descent"', '—'],
        ['2', 'Sequential, Dense', '3B1B "Neural Networks" Ch.1', '퍼셉트론 논문 리뷰'],
        ['3', 'EarlyStopping, Sigmoid', '3B1B "Backpropagation"', '과적합 블로그'],
        ['4', 'Dropout, save/load', 'StatQuest "Dropout"', 'Dropout 논문'],
        ['5', 'Conv2D, MaxPool', '3B1B "Convolution"', 'AlexNet 논문 리뷰'],
        ['6', 'RNN, LSTM', '3B1B "RNN"', 'LSTM 논문 리뷰'],
        ['7', 'Embedding, Tokenizer', '3B1B "Word2Vec"', 'Word2Vec 논문 리뷰'],
        ['8', 'VGG16, ResNet', '구글 "VGGNet 논문 리뷰"', 'VGG/ResNet 논문'],
        ['9', 'PCA, KFold, XGBoost', '3B1B "Linear Algebra"', 'XGBoost 논문 리뷰'],
        ['10', 'Transformer, Attention', '3B1B "Transformers"', 'Attention Is All You Need'],
    ]
)

# ══════════════════════════════════════════
# 8. 세계 트렌드와 HEART Lab
# ══════════════════════════════════════════
doc.add_heading('7. 세계 트렌드와 HEART Lab 연결', level=1)

doc.add_heading('글로벌 트렌드', level=2)
add_table(doc,
    ['분야', '키워드', '대표'],
    [
        ['감정 AI', 'Emotion AI, Affective Computing', 'NVIDIA Audio2Face, Affectiva'],
        ['LLM', '대규모 언어모델', 'GPT, Claude, Gemini'],
        ['AI Agent', '자율적 의사결정', '자율주행, 로봇 제어'],
        ['멀티모달', '영상 + 음성 + 텍스트', 'GPT-4V, Gemini'],
        ['Physical AI', '로봇, 자율주행', 'NVIDIA Isaac, Figure AI'],
    ]
)

doc.add_heading('HEART Lab 과제', level=2)
add_table(doc,
    ['과제', '파트너', '내용', '커리큘럼 연결'],
    [
        ['차량 감정인식', '현대자동차', '멀티모달 센싱 → 감정 → 안전운전', 'Phase 5~7'],
        ['협동로봇 AI', '두산', 'VLM/VLA + AI SoC + 강화학습', 'Phase 8~10'],
        ['디지털 트윈', '토론토대', 'EV 배터리 + AI Agent', 'Phase 10'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 9. 폴더 구조
# ══════════════════════════════════════════
doc.add_heading('8. 폴더 구조', level=1)

folder_text = """HEART-Lab-Curriculum/
├── README.md
├── CURRICULUM_MASTER.md
├── DIAGNOSTIC_TEST.md
├── MANUAL.md
│
├── Phase00_why/
├── Phase01_foundations/
├── Phase02_get_started/
├── Phase03_debugging/
├── Phase04_save_structure/
├── Phase05_image/
├── Phase06_timeseries/
├── Phase07_nlp/
├── Phase08_optimization/
├── Phase09_ml/
├── Phase10_pytorch/
│
├── worksheets/
│   ├── checkpoint_1.md
│   ├── checkpoint_2.md
│   └── final_project.md
│
└── AI5-main/  (원본 학습 코드)"""
add_code_block(doc, folder_text)

# ══════════════════════════════════════════
# 마무리
# ══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '#0→#1→#2→#3→#4 뼈대를 몸에 익히고,\n'
    '한 번에 하나씩만 바꿔가며,\n'
    '같은 데이터에 다른 기법을 적용하면서,\n'
    '3Blue1Brown과 논문 리뷰로 "왜"를 이해하고,\n'
    'HEART Lab의 실제 과제에 투입될 수 있는 연구원을 만든다.'
)
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ── 저장 ──
output_path = os.path.join(os.path.dirname(__file__), 'CURRICULUM_MASTER.docx')
doc.save(output_path)
print(f'저장 완료: {output_path}')
